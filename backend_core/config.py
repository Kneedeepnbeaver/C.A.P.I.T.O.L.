import sys
from pathlib import Path
import json

# --- CONFIGURATION ---
if getattr(sys, 'frozen', False):
    # PyInstaller creates a temp folder and stores path in _MEIPASS, 
    # but for persistent data, we want the directory of the executable
    BASE_DIR = Path(sys.executable).parent
    # If using --onefile, sys.executable is the binary, so parent is the folder containing it.
else:
    # In development, use project root (assuming this file is in backend_core/)
    BASE_DIR = Path(__file__).parent.parent.resolve()

METADATA_CSV = BASE_DIR / "legislative_documents" / "metadata.csv" # Adjusted path if needed, or keep same?
# Original was: METADATA_CSV = BASE_DIR / "legislative_metadata.csv"
# But server_rag.py often used a different path? 
# Let's clean this up. We will stick to the original "legislative_metadata.csv" in BASE_DIR for now to match legacy.
METADATA_CSV = BASE_DIR / "legislative_metadata.csv"

TEXT_DIR = BASE_DIR / "legislative_documents"
OUTPUT_DIR = BASE_DIR / "generated_content"
MODEL = "llama3.2:latest"
VOICE_PRESETS_FILE = BASE_DIR / "legislative_voice_presets.json"

# Ensure directories exist
TEXT_DIR.mkdir(exist_ok=True, parents=True)
OUTPUT_DIR.mkdir(exist_ok=True, parents=True)

def load_voice_presets():
    try:
        if VOICE_PRESETS_FILE.exists():
            with open(VOICE_PRESETS_FILE, 'r') as f:
                return json.load(f)['voice_presets']
    except Exception as e:
        print(f"Error loading presets: {e}")
    return []

# --- UTILITY FUNCTIONS ---
# Originally in legislative_backend.py, moved here for server_rag compatibility

def list_ollama_models():
    """Returns a list of installed Ollama model names."""
    import subprocess
    
    # Packaged apps on macOS don't inherit the user's PATH
    # We must try common installation locations
    ollama_paths = ["ollama", "/usr/local/bin/ollama", "/opt/homebrew/bin/ollama"]
    
    for cmd in ollama_paths:
        try:
            result = subprocess.run(
                [cmd, "list"],
                capture_output=True,
                text=True
            )
            if result.returncode == 0:
                lines = result.stdout.strip().split('\n')
                if len(lines) <= 1: return []
                return [line.split()[0] for line in lines[1:] if line.strip()]
        except FileNotFoundError:
            continue
        except Exception as e:
            print(f"Error checking ollama at {cmd}: {e}")
            continue
            
    return []

def get_full_text(filename):
    """Reads the full text of a document from the text directory."""
    import pdfplumber
    path = TEXT_DIR / filename
    if not path.exists():
        return ""
        
    try:
        suffix = path.suffix.lower()
        if suffix == '.pdf':
            text = ""
            try:
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                return text
            except Exception as e:
                return f"[Error reading PDF: {e}]"
        
        elif suffix == '.docx':
            try:
                from docx import Document
                doc = Document(path)
                return "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                return f"[Error reading DOCX: {e}]"
        
        else:
            # Assume Text/Markdown/CSV etc
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(path, 'r', encoding='latin-1') as f:
                    return f.read()
    except Exception as e:
        return f"[Error reading file: {e}]"

def search_documents(query):
    """
    Searches for the query string in the full text of all documents.
    Returns a list of matching filenames.
    """
    query = query.lower()
    results = []
    
    # Iterate over all files in the text directory
    if not TEXT_DIR.exists():
        return []
        
    for file_path in TEXT_DIR.iterdir():
        if file_path.is_file() and file_path.name != '.DS_Store':
            # We use the existing get_full_text to safely handle PDFs/encoding
            text = get_full_text(file_path.name)
            if query in text.lower():
                results.append(file_path.name)
                
    return results

def call_ollama(prompt, model=None):
    """Call Ollama via HTTP API (avoids subprocess deadlock issues)."""
    import requests
    
    target_model = model if model else MODEL
    try:
        response = requests.post(
            'http://localhost:11434/api/generate',
            json={
                'model': target_model,
                'prompt': prompt,
                'stream': False
            },
            timeout=180  # 3 minutes
        )
        
        if response.status_code == 200:
            return response.json().get('response', '')
        else:
            return f"Error: Ollama returned status {response.status_code}"
            
    except requests.exceptions.Timeout:
        return "Error: Generation timed out. Try phi3:latest for faster results."
    except requests.exceptions.ConnectionError:
        return "Error: Could not connect to Ollama. Make sure it's running."

def generate_legislative_artifact(selected_docs, artifact_type, tone, additional_instructions, model=None):
    """
    OPTIMIZED: Generates artifacts using ONLY pre-extracted metadata.
    Enhanced prompts for better quality output.
    """
    
    # Build context using ONLY metadata (no file reading!)
    context = ""
    for i, doc in enumerate(selected_docs):
        context += f"═══ DOCUMENT {i+1} ═══\n"
        context += f"Title: {doc.get('Document Title', 'Untitled')}\n"
        context += f"Sender: {doc.get('Sender/Organization', 'Unknown')}\n"
        context += f"Position: {doc.get('Position', 'Neutral')}\n"
        context += f"Bill: {doc.get('Bill Number', 'N/A')}\n\n"
        
        # Use rich metadata fields (already extracted by AI during import)
        if doc.get('Summary'):
            context += f"SUMMARY:\n{doc['Summary']}\n\n"
        
        if doc.get('Key Arguments'):
            context += f"KEY ARGUMENTS:\n{doc['Key Arguments']}\n\n"
        
        if doc.get('Stakeholders'):
            context += f"STAKEHOLDERS:\n{doc['Stakeholders']}\n\n"
        
        if doc.get('Keywords'):
            context += f"KEYWORDS: {doc['Keywords']}\n\n"
        
        context += "\n"
    
    # Enhanced prompts with structure and examples
    prompts = {
        "Executive Summary": """Create a professional Executive Summary with this structure:
1. **Issue Overview**: Core legislative issue and context
2. **Stakeholder Positions**: Who supports/opposes and their key reasons
3. **Critical Arguments**: Main points from both sides
4. **Legislative Landscape**: Political context and likely outcome

Use clear headers and be objective.""",

        "Talking Points (Pro)": """Create persuasive Talking Points SUPPORTING the legislation:
- Lead with the strongest argument
- Use concrete examples from the documents
- Frame benefits for constituents
- Address objections preemptively
- Keep each point to 1-2 sentences""",

        "Talking Points (Con)": """Create persuasive Talking Points OPPOSING the legislation:
- Lead with the most serious concern
- Cite specific problems from testimony
- Highlight costs or unintended consequences
- Use concrete examples
- Keep each point to 1-2 sentences""",

        "Vote Recommendation": """Write a Vote Recommendation Memo with:
1. **Recommendation**: Clear AYE or NO with rationale
2. **Policy Analysis**: Merits and concerns
3. **Political Considerations**: Constituent impact
4. **Risk Assessment**: Consequences of each vote
5. **Conclusion**: Restate recommendation""",

        "Coalition Letter": """Draft a Coalition Letter with:
- Clear call to action (Vote YES/NO)
- Unified voice for all coalition members
- 2-3 strongest shared arguments
- Specific request and deadline
- Professional format""",

        "Opposition Research": """Create Opposition Research identifying:
- Contradictions in arguments
- Weak or unsupported claims
- Controversial statements
- Vulnerabilities to exploit
- Vulnerabilities to exploit
Use direct quotes where possible."""
    }
    
    # New additions
    prompts.update({
        "Policy Analysis": """Perform a deep-dive Policy Analysis covering:
1. **Legal & Regulatory Framework**: How this interacts with existing law.
2. **Economic & Social Impact**: Quantifiable and qualitative consequences.
3. **Implementation Challenges**: Practical hurdles for agencies.
4. **Conclusion**: Overall merit of the policy.""",

        "Policy Recommendations": """Generate forward-looking Policy Recommendations:
1. **Strategic Adjustments**: How to improve the current bill/policy.
2. **Future Legislation**: Related issues that will need addressing next session.
3. **Best Practices**: Model language from other jurisdictions.
4. **Action Plan**: Immediate next steps for proponents.""",

        "Press Release": """Draft a professional Press Release:
- **Header**: Immediate Release info and catchy headline.
- **Lead**: The 'Who, What, Where, When, Why'.
- **Body**: Core details and significance.
- **Quote Placeholder**: Compelling quote from a key stakeholder.
- **Boilerplate**: Standard background info.""",

        "Social Media Suite": """Create a Social Media Content Suite:
- **X/Twitter**: 3-5 thread-ready posts with hashtags.
- **LinkedIn**: A long-form thought leadership post.
- **Instagram/Facebook**: Concise, high-impact captions focusing on the 'human' element.
- **Tone**: Engaging, sharable, and informative.""",

        "Committee Briefing": """Draft a condensed Committee Briefing:
1. **The 'Bottom Line'**: 30-second elevator pitch.
2. **Key Questions**: Hard questions to ask witnesses.
3. **Political Temperature**: Who is leaning which way.
4. **Critical Timing**: Deadlines and procedural notes."""
    })

    base_instruction = prompts.get(artifact_type, "Analyze and synthesize the documents.")
    
    # Enhanced system prompt
    final_prompt = f"""You are a senior legislative analyst with expertise in California state politics, policy analysis, and political strategy.

TASK: {artifact_type}

INSTRUCTIONS:
{base_instruction}

TONE: {tone}
{f"ADDITIONAL: {additional_instructions}" if additional_instructions else ""}

SOURCE DOCUMENTS:
{context}

FORMAT REQUIREMENTS:
- Use Markdown with clear headers (##, ###)
- Cite specific organizations and arguments
- Be actionable and concrete
- Start with: # {artifact_type}

Begin:"""

    return call_ollama(final_prompt, model=model)

# --- EXPORT LOGIC ---

HTML_STYLES = {
    "Professional": """
        body { font-family: 'Helvetica Neue', Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 20px; color: #333; }
        h1 { color: #2c3e50; border-bottom: 2px solid #2c3e50; padding-bottom: 10px; }
        h2 { color: #34495e; margin-top: 30px; }
        ul { margin-bottom: 20px; }
        li { margin-bottom: 5px; }
    """,
    "Academic": """
        body { font-family: 'Georgia', serif; line-height: 1.8; max-width: 800px; margin: 40px auto; padding: 20px; color: #111; }
        h1 { font-family: 'Arial', sans-serif; text-align: center; }
        h2 { font-style: italic; border-bottom: 1px solid #ccc; }
    """,
    "Draft": """
        body { font-family: 'Courier New', monospace; line-height: 1.4; max-width: 800px; margin: 40px auto; padding: 20px; color: #555; }
    """
}

def export_to_file(content, format_type, style_name="Professional", filename_prefix="leg_analysis"):
    """
    Exports content to a file.
    format_type: 'markdown', 'html', 'docx', 'txt'
    Returns the absolute path of the created file or None if error.
    """
    from datetime import datetime
    import markdown
    from docx import Document

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if format_type.lower() in ["markdown", "md"]:
        filename = f"{filename_prefix}_{timestamp}.md"
        filepath = OUTPUT_DIR / filename
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return str(filepath)
        
    elif format_type.lower() == "txt":
        filename = f"{filename_prefix}_{timestamp}.txt"
        filepath = OUTPUT_DIR / filename
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return str(filepath)
        
    elif format_type.lower() == "html":
        filename = f"{filename_prefix}_{timestamp}.html"
        filepath = OUTPUT_DIR / filename
        
        # Convert MD to HTML
        html_body = markdown.markdown(content)
        css = HTML_STYLES.get(style_name, HTML_STYLES["Professional"])
        
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>{css}</style>
</head>
<body>
    {html_body}
</body>
</html>"""
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(full_html)
        return str(filepath)
        
    elif format_type.lower() == "docx":
        filename = f"{filename_prefix}_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename
        
        doc = Document()
        # Simple Markdown parsing
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        doc.save(filepath)
        return str(filepath)
        
    return None
